from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
FIGURES = ROOT / "figures"
OUTPUT = ROOT / "Munir_Khaliqyar_MiniPhD_Final_Paper_v2.docx"
EXAMPLES_DIR = ROOT / "llm_multiturn_attacks-main" / "llm_multiturn_attacks-main" / "examples"
CHAT_PATH = ROOT / "ChatAlpaca-main" / "ChatAlpaca-main" / "data" / "chatalpaca-10k.json"
CONTURE_PATH = ROOT / "conture-main" / "conture-main" / "data" / "data.json"


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)


def load_prompt_level_results() -> pd.DataFrame:
    rows = []
    for csv_path in EXAMPLES_DIR.rglob("*.csv"):
        df = pd.read_csv(csv_path)
        parts = csv_path.stem.split("_")
        turns = None
        for idx, part in enumerate(parts):
            if part == "N" and idx + 1 < len(parts):
                turns = int(parts[idx + 1])
                break
        tense = parts[-1]
        model = csv_path.parent.name
        final_df = df if turns == 1 else df[df["Multi Step"] == turns]
        tmp = final_df[["Unique ID", "Human"]].copy()
        tmp["Human"] = pd.to_numeric(tmp["Human"], errors="coerce").fillna(0).astype(int)
        tmp["model"] = model
        tmp["tense"] = tense
        tmp["turns"] = turns
        rows.append(tmp)

    all_df = pd.concat(rows, ignore_index=True)
    wide = all_df.pivot_table(
        index=["model", "Unique ID"], columns=["tense", "turns"], values="Human", aggfunc="first"
    ).reset_index()
    wide.columns = ["model", "Unique ID", "past_1", "past_2", "past_3", "present_1", "present_2", "present_3"]
    wide["any_single_turn_warning"] = ((wide["present_1"] == 1) | (wide["past_1"] == 1)).astype(int)
    wide["any_three_turn_jailbreak"] = ((wide["present_3"] == 1) | (wide["past_3"] == 1)).astype(int)
    wide["hidden_present_failure"] = ((wide["present_1"] == 0) & (wide["present_3"] == 1)).astype(int)
    return wide


def metrics(df: pd.DataFrame, predictor: str, outcome: str) -> tuple[float, float]:
    tp = ((df[predictor] == 1) & (df[outcome] == 1)).sum()
    fp = ((df[predictor] == 1) & (df[outcome] == 0)).sum()
    fn = ((df[predictor] == 0) & (df[outcome] == 1)).sum()
    precision = tp / (tp + fp) if tp + fp else 0.0
    recall = tp / (tp + fn) if tp + fn else 0.0
    return precision, recall


def compute_chat_stats() -> tuple[float, float]:
    import json

    conversations = []
    for line in CHAT_PATH.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line:
            conversations.append(json.loads(line)["conversations"])
    avg_messages = sum(len(c) for c in conversations) / len(conversations)
    avg_user_turns = sum(sum(1 for t in c if t["from"] == "human") for c in conversations) / len(conversations)
    return avg_messages, avg_user_turns


def compute_conture_stats() -> tuple[float, float]:
    import json

    dialogs = json.loads(CONTURE_PATH.read_text(encoding="utf-8"))
    turn1 = []
    turn9 = []
    for dialog in dialogs:
        if dialog["turns"]:
            turn1.append(dialog["turns"][0]["overall impression"])
        if len(dialog["turns"]) >= 9:
            turn9.append(dialog["turns"][8]["overall impression"])
    return sum(turn1) / len(turn1), sum(turn9) / len(turn9)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Do Single-Turn Safety Benchmarks Actually Predict Multi-Turn Jailbreak Vulnerability?")
    run.bold = True
    run.font.size = Pt(15)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Munir Khaliqyar\nDepartment of Computer Science, Bishop's University, Sherbrooke, Quebec, Canada")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, italic: bool = False) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.italic = italic


def add_figure(document: Document, image_name: str, caption: str, width: float = 5.9) -> None:
    document.add_picture(str(FIGURES / image_name), width=Inches(width))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True


def build_document() -> Document:
    doc = Document()
    set_default_font(doc)
    add_title(doc)

    wide = load_prompt_level_results()
    hidden_rate = wide["hidden_present_failure"].mean() * 100.0
    p_direct, r_direct = metrics(wide, "present_1", "any_three_turn_jailbreak")
    p_reframed, r_reframed = metrics(wide, "past_1", "any_three_turn_jailbreak")
    p_either, r_either = metrics(wide, "any_single_turn_warning", "any_three_turn_jailbreak")
    avg_messages, avg_user_turns = compute_chat_stats()
    conture_turn1, conture_turn9 = compute_conture_stats()
    present_1 = wide["present_1"].mean() * 100.0
    present_3 = wide["present_3"].mean() * 100.0
    past_1 = wide["past_1"].mean() * 100.0
    past_3 = wide["past_3"].mean() * 100.0

    add_heading(doc, "Abstract")
    add_paragraph(
        doc,
        "Single-turn safety evaluations remain attractive because they are cheap, standardized, and easy to score, but they can understate how chatbots behave in real conversations. This paper asks whether single-turn safety benchmarks actually predict multi-turn jailbreak vulnerability. Using three local datasets, I argue that ordinary direct single-turn benchmarks predict later conversational failure poorly, but a richer low-cost screen does better when it measures three properties: retryability, decomposability, and benign-indistinguishability. The main evidence comes from a jailbreak dataset containing 127 harmful prompt families tested across four models under direct single-turn, reframed single-turn, and decomposed multi-turn variants. Direct single-turn evaluation misses many failures: 24.6% of prompt-model pairs are hidden failures that appear safe in a direct one-turn test but jailbreak by turn three. A cheap stressed single-turn warning signal improves recall for later three-turn jailbreaks from 18.8% to 39.7% while maintaining 85.4% precision. Supporting evidence from ChatAlpaca and ConTurE shows that chatbot interactions are naturally multi-turn and that response behavior changes across turns. The paper argues that direct single-turn benchmarks are weak predictors of conversational jailbreak risk, while a three-part low-cost screen is a more useful early-warning proxy."
    )

    add_heading(doc, "1. Introduction")
    add_paragraph(
        doc,
        "Large language model safety evaluation still relies heavily on single-turn prompting, even though real misuse often unfolds through gradual escalation, reframing, persistence, and context accumulation [1, 2, 3, 4]. This creates a benchmark-validity problem: a model can appear safe when evaluated on isolated prompts but fail in realistic dialogue. The practical question is not only whether multi-turn attacks are stronger than single-turn attacks, but whether inexpensive tests can predict which apparently safe systems are most likely to fail later."
    )
    add_paragraph(
        doc,
        "Recent work on jailbreaks has shown that aligned models remain vulnerable to adversarial prompting, transfer attacks, in-the-wild jailbreak prompts, and automated red-teaming pipelines [5, 6, 7, 8, 9]. Multi-turn attacks are especially important because they exploit the conversational setting that many public chatbots are designed for [3, 10]. In parallel, broader evaluation work has emphasized that trustworthiness and safety cannot be inferred from a single narrow benchmark alone [11, 12]. NIST guidance similarly treats risk as a system-level issue that depends on context, deployment, and evaluation design rather than a single static score [13]."
    )
    add_paragraph(
        doc,
        "This mini-PhD addresses the following research question: Do single-turn safety benchmarks actually predict multi-turn jailbreak vulnerability? My thesis is that direct single-turn benchmarks are too weak because they mostly test direct refusal, while multi-turn vulnerability depends on retryability, decomposability, and benign-indistinguishability. A low-cost evaluation that measures those three properties provides a better early-warning signal of later conversational jailbreak risk."
    )

    add_heading(doc, "2. Methods and Data")
    add_paragraph(
        doc,
        "The main analysis uses the local repository llm_multiturn_attacks-main, which contains 127 cybersecurity-related harmful prompt families derived from adversarial benchmark sources, along with two-turn and three-turn decompositions and example jailbreak outcomes for four target models: GPT-4o-mini, Gemini-2.0-Flash, Llama-2-7b-chat, and Qwen2-7b-instruct [3, 14, 15]. Human labels indicate whether the final response was successfully jailbroken."
    )
    add_paragraph(
        doc,
        "I constructed prompt-level comparisons by aligning the same prompt family across direct present-tense single-turn evaluation, past-tense reframed single-turn evaluation, and final-turn outcomes from corresponding two-turn and three-turn attacks. The core prompt-model table contains 508 prompt-model pairs. The primary outcome is whether the three-turn attack produces a human-labeled jailbreak. The main predictor variables are direct single-turn jailbreak, reframed single-turn jailbreak, and a combined warning signal that fires when either single-turn condition jailbreaks."
    )
    add_paragraph(
        doc,
        "I report four simple statistics: success rates, a hidden-failure rate, conditional multi-turn risk, and predictor precision and recall. A hidden failure is defined as a prompt-model pair that is safe on the direct present-tense single-turn test but unsafe on the final turn of the corresponding three-turn present-tense attack. Precision measures how often a cheap warning signal correctly identifies future three-turn jailbreaks, while recall measures how many later jailbreaks are caught by that signal."
    )
    add_paragraph(
        doc,
        "The proposed evaluation framework has three dimensions. Retryability asks whether a harmful goal is already fragile enough to succeed under small single-turn changes such as paraphrase, retry, or historical reframing. Decomposability asks whether the same harmful goal can be spread across multiple individually less suspicious turns. Benign-indistinguishability asks whether those turns can remain superficially educational, descriptive, or otherwise normal-looking while still advancing the harmful objective. The current local data measures retryability and decomposability more directly than benign-indistinguishability, so the third dimension should be treated as a conceptually important but only partially operationalized part of the framework."
    )
    add_paragraph(
        doc,
        f"Two supporting datasets provide external context. ChatAlpaca contains 10,000 generated multi-turn conversations, with an average of {avg_messages:.2f} total messages and {avg_user_turns:.2f} user turns per conversation [16]. ConTurE contains 119 human-annotated dialogs with turn-level quality ratings; in the local copy, mean quality falls from {conture_turn1:.2f} on turn 1 to {conture_turn9:.2f} on turn 9 [17, 18]. These datasets do not measure jailbreaks directly, but they support the argument that chatbot behavior is inherently multi-turn and dynamic."
    )

    add_heading(doc, "3. Results")
    add_paragraph(
        doc,
        f"Multi-turn escalation substantially increases jailbreak success. In the local attack results, mean jailbreak success rises from {present_1:.1f}% to {present_3:.1f}% for present-tense prompts and from {past_1:.1f}% to {past_3:.1f}% for past-tense prompts when moving from one turn to three turns. This confirms that both conversational decomposition and seemingly mild reframing increase risk."
    )
    add_paragraph(
        doc,
        f"The most important finding is that direct single-turn testing misses many later failures. Across the 508 prompt-model pairs, {hidden_rate:.1f}% are hidden failures: they appear safe in a direct one-turn test but jailbreak by the third turn. This means that a direct single-turn benchmark can produce substantial false confidence even before we compare across different models."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Predictor"
    hdr[1].text = "Precision"
    hdr[2].text = "Recall"
    rows = [
        ("Direct single-turn warning", f"{p_direct * 100:.1f}%", f"{r_direct * 100:.1f}%"),
        ("Reframed single-turn warning", f"{p_reframed * 100:.1f}%", f"{r_reframed * 100:.1f}%"),
        ("Either single-turn warning", f"{p_either * 100:.1f}%", f"{r_either * 100:.1f}%"),
    ]
    for predictor, precision, recall in rows:
        cells = table.add_row().cells
        cells[0].text = predictor
        cells[1].text = precision
        cells[2].text = recall

    add_paragraph(
        doc,
        "These results show that direct single-turn evaluation is a high-precision but low-recall predictor. When a direct one-turn prompt already jailbreaks, the system is often genuinely risky, but this signal misses most future three-turn failures. Reframed single-turn evaluation catches more future jailbreaks, and the combined warning signal performs best overall as a low-cost screening proxy. Interpreted through the proposed framework, retryability explains why reframed single-turn prompts reveal near-boundary fragility, while decomposability explains why success grows once the harmful request is split across turns. Benign-indistinguishability explains why these multi-turn sequences can survive longer than a direct harmful request: the early turns do not always look obviously unsafe."
    )

    add_figure(
        doc,
        "hidden_present_failures.png",
        "Figure 1. Hidden failures by model: prompt-model pairs that look safe in direct single-turn evaluation but jailbreak by turn three.",
    )
    add_figure(
        doc,
        "predictor_precision_recall.png",
        "Figure 2. Precision and recall for low-cost single-turn predictors of later three-turn jailbreaks.",
    )
    add_figure(
        doc,
        "conditional_multi_turn_risk.png",
        "Figure 3. Three-turn jailbreak risk is substantially higher once the cheap single-turn screen already raises a warning.",
    )
    add_figure(
        doc,
        "chatalpaca_turn_distribution.png",
        "Figure 4. ChatAlpaca shows that chatbot interactions are naturally multi-turn rather than one-shot.",
    )

    add_heading(doc, "4. Discussion")
    add_paragraph(
        doc,
        "The central contribution of this project is a change in framing. Instead of asking only whether single-turn and multi-turn benchmarks differ, I ask whether single-turn benchmarks actually predict multi-turn jailbreak vulnerability and, if not, what low-cost signal does better. The answer from the local data is that ordinary direct single-turn benchmarks predict poorly because they ignore retryability, decomposability, and benign-indistinguishability. A lightly stressed single-turn test performs better because it pushes prompts closer to the refusal boundary without requiring a full multi-turn red-team run."
    )
    add_paragraph(
        doc,
        "This suggests a staged evaluation workflow for AI safety practice. First, run a cheap single-turn screen that includes direct prompts plus a small family of reframed variants to estimate retryability. Second, test a short decomposed conversation to estimate decomposability. Third, assess whether the intermediate steps remain superficially benign enough to bypass simple filters. The approach does not replace conversational evaluation, but it can make safety testing more scalable by focusing attention where hidden risk is most likely."
    )
    add_paragraph(
        doc,
        "There are also clear limits. The current evidence is concentrated in a cybersecurity jailbreak dataset, and the strongest predictor is still only a screening signal. Even the best combined warning signal reaches 39.7% recall, which means many conversational jailbreaks remain invisible unless the evaluator explicitly tests multi-turn behavior. In addition, benign-indistinguishability is better motivated than directly quantified in the current paper. For that reason, the strongest defensible claim is not that single-turn evaluation is sufficient, but that a three-part low-cost screen is a better early-warning method than direct one-turn benchmarking alone."
    )

    add_heading(doc, "5. Limitations")
    add_paragraph(
        doc,
        "The paper has four main limitations. First, the main safety evidence comes from one domain-specific jailbreak dataset rather than a broad cross-domain benchmark. Second, the predictor used here is intentionally simple and hand-designed; a richer classifier built from prompt features could improve performance. Third, ChatAlpaca and ConTurE support the case for conversational evaluation but do not directly measure harmful-output safety. Fourth, benign-indistinguishability is only partially operationalized in the current dataset, so that part of the framework remains more conceptual than fully validated."
    )

    add_heading(doc, "6. Conclusion")
    add_paragraph(
        doc,
        "Single-turn safety benchmarks can create false confidence because many failures only emerge after reframing and context accumulation. In the local data, nearly one quarter of prompt-model pairs are hidden failures that appear safe on direct single-turn evaluation but jailbreak by turn three. The broader explanation is that direct single-turn tests fail to measure retryability, decomposability, and benign-indistinguishability together. A low-cost screen that moves closer to those dimensions improves prediction of later multi-turn failure while preserving high precision. The right conclusion is therefore not that multi-turn testing is optional, but that better low-cost screening can help identify where full conversational testing is most urgently needed."
    )

    add_heading(doc, "References")
    refs = [
        "[1] Deep Ganguli et al. Red Teaming Language Models to Reduce Harms: Methods, Scaling Behaviors, and Lessons Learned. arXiv:2209.07858, 2022.",
        "[2] Alexander Wei, Nika Haghtalab, and Jacob Steinhardt. Jailbroken: How Does LLM Safety Training Fail? arXiv:2307.02483, 2023.",
        "[3] Michael Tchuindjang, Nathan Duran, Phil Legg, and Faiza Medjek. Jailbreaking LLMs Through Tense Manipulation in Multi-turn Dialogues. In Advances in Computational Intelligence Systems, Springer, 2026.",
        "[4] NIST. Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile. NIST AI 600-1, 2024.",
        "[5] Andy Zou, Zifan Wang, Nicholas Carlini, Milad Nasr, J. Zico Kolter, and Matt Fredrikson. Universal and Transferable Adversarial Attacks on Aligned Language Models. arXiv:2307.15043, 2023.",
        "[6] Xinyue Shen, Zeyuan Chen, Michael Backes, Yun Shen, and Yang Zhang. Do Anything Now: Characterizing and Evaluating In-The-Wild Jailbreak Prompts on Large Language Models. arXiv:2308.03825, 2024.",
        "[7] Mantas Mazeika et al. HarmBench: A Standardized Evaluation Framework for Automated Red Teaming and Robust Refusal. arXiv:2402.04249, 2024.",
        "[8] Boxin Wang et al. DecodingTrust: A Comprehensive Assessment of Trustworthiness in GPT Models. arXiv:2306.11698, 2023.",
        "[9] Annie Russinovich et al. Crescendo: An Automatic Multi-Turn LLM Jailbreak Attack for Current and Future LLMs. 2024.",
        "[10] Xiaogeng Liu, Nan Xu, Muhao Chen, and Chaowei Xiao. AutoDAN: Generating Stealthy Jailbreak Prompts on Aligned Large Language Models. arXiv:2310.04451, 2023.",
        "[11] Laura Weidinger et al. Taxonomy of risks posed by language models. FAccT, 2022.",
        "[12] Percy Liang et al. Holistic Evaluation of Language Models. arXiv:2211.09110, 2022.",
        "[13] NIST. AI Risk Management Framework 1.0. 2023.",
        "[14] AdvBench benchmark repository: llm-attacks/llm-attacks. 2023.",
        "[15] Center for AI Safety. HarmBench benchmark repository. 2024.",
        "[16] icip-cas. ChatAlpaca: A Multi-Turn Dialogue Corpus based on Alpaca Instructions. GitHub repository, 2023.",
        "[17] Sarik Ghazarian, Behnam Hedayatnia, Alexandros Papangelis, Yang Liu, and Dilek Hakkani-Tur. What is wrong with you?: Leveraging User Sentiment for Automatic Dialog Evaluation. Findings of ACL, 2022.",
        "[18] Chulaka Gunasekara et al. Overview of the Ninth Dialog System Technology Challenge: DSTC9. AAAI Workshop Proceedings, 2021.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
